from fpdf import FPDF
from docx import Document
import json
import os
import re

FONT_REGULAR = os.path.join(os.path.dirname(__file__), "DejaVuSans.ttf")
FONT_BOLD = os.path.join(os.path.dirname(__file__), "DejaVuSans-Bold.ttf")

def parse_protocol_sections(text: str):
    """Flexible parser for structured output."""
    sections = []
    current_title = None
    current_body = []

    for line in text.splitlines():
        line = line.strip()
        is_section_header = (
            re.match(r"^#{2,}\s+", line) or
            (line.isupper() and len(line) < 60) or
            re.match(r"^[A-Za-z\s]+:$", line)
        )
        if is_section_header:
            if current_title and current_body:
                sections.append((current_title, "\n".join(current_body).strip()))
                current_body = []
            clean_title = re.sub(r"^#+\s*", "", line).replace(":", "").strip()
            current_title = clean_title
        else:
            current_body.append(line)
    if current_title and current_body:
        sections.append((current_title, "\n".join(current_body).strip()))
    return sections

def save_structured_protocol_as_docx(protocol: dict, filename: str = "structured_protocol.docx") -> str:
    doc = Document()
    
    def add_heading(title):
        doc.add_heading(title, level=2)
    
    def add_paragraph(text):
        for line in text.split('\n'):
            doc.add_paragraph(line.strip())

    add_heading("Title")
    add_paragraph(protocol["title"])
    
    add_heading("Background")
    add_paragraph(protocol["background"])

    add_heading("Rationale")
    add_paragraph(protocol["rationale"])

    add_heading("Objectives")
    doc.add_paragraph("Primary Objective", style="List Bullet")
    add_paragraph(protocol["objectives"]["primary"])
    doc.add_paragraph("Secondary Objectives", style="List Bullet")
    add_paragraph(protocol["objectives"]["secondary"])

    add_heading("Study Design")
    add_paragraph(protocol["study_design"])

    add_heading("Population")
    doc.add_paragraph("Inclusion Criteria", style="List Bullet")
    add_paragraph(protocol["population"]["inclusion_criteria"])
    doc.add_paragraph("Exclusion Criteria", style="List Bullet")
    add_paragraph(protocol["population"]["exclusion_criteria"])

    add_heading("Intervention")
    add_paragraph(protocol["intervention"])

    add_heading("Endpoints")
    doc.add_paragraph("Primary Endpoint", style="List Bullet")
    add_paragraph(protocol["endpoints"]["primary"])
    doc.add_paragraph("Secondary Endpoints", style="List Bullet")
    add_paragraph(protocol["endpoints"]["secondary"])

    add_heading("Safety Monitoring")
    add_paragraph(protocol["safety_monitoring"])

    add_heading("Statistical Analysis")
    add_paragraph(protocol["statistical_analysis"])

    add_heading("Ethical Considerations")
    add_paragraph(protocol["ethical_considerations"])

    path = f"/tmp/{filename}"
    doc.save(path)
    return path

class UnicodePDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_page()
        self.set_auto_page_break(auto=True, margin=15)
        self.set_left_margin(10)
        self.set_right_margin(10)
        self.add_font("DejaVu", "", FONT_REGULAR, uni=True)
        self.add_font("DejaVu", "B", FONT_BOLD, uni=True)
        self.set_font("DejaVu", size=12)

def save_structured_protocol_as_pdf(protocol: dict, filename: str = "structured_protocol.pdf") -> str:
    pdf = UnicodePDF()

    def add_section(title, content):
        pdf.set_font("DejaVu", "B", 14)
        pdf.multi_cell(0, 10, title)
        pdf.set_font("DejaVu", "", 12)

        # Break large lines safely
        for line in content.split("\n"):
            # If any line is too long without spaces, manually insert breaks
            safe_line = re.sub(r'([^\s]{60})(?=[^\s])', r'\1\n', line)
            try:
                pdf.multi_cell(0, 8, safe_line)
            except Exception as e:
                pdf.multi_cell(0, 8, "[Line could not be rendered]")
                print(f"[PDF Line Error] in section '{title}': {e}")

        pdf.ln(4)


    add_section("Title", protocol["title"])
    add_section("Background", protocol["background"])
    add_section("Rationale", protocol["rationale"])
    add_section("Primary Objective", protocol["objectives"]["primary"])
    add_section("Secondary Objectives", protocol["objectives"]["secondary"])
    add_section("Study Design", protocol["study_design"])
    add_section("Inclusion Criteria", protocol["population"]["inclusion_criteria"])
    add_section("Exclusion Criteria", protocol["population"]["exclusion_criteria"])
    add_section("Intervention", protocol["intervention"])
    add_section("Primary Endpoint", protocol["endpoints"]["primary"])
    add_section("Secondary Endpoints", protocol["endpoints"]["secondary"])
    add_section("Safety Monitoring", protocol["safety_monitoring"])
    add_section("Statistical Analysis", protocol["statistical_analysis"])
    add_section("Ethical Considerations", protocol["ethical_considerations"])

    path = f"/tmp/{filename}"
    pdf.output(path)
    return path

def save_protocol_as_json(user_inputs: dict, protocol: dict, filename: str = "structured_protocol.json") -> str:
    """
    Save the structured protocol as JSON, including both metadata (inputs) and full protocol content.
    """
    json_obj = {
        "metadata": {
            "condition": user_inputs.get("condition"),
            "intervention": user_inputs.get("intervention"),
            "population": user_inputs.get("population")
        },
        "protocol": {
            "title": protocol["title"],
            "background": protocol["background"],
            "rationale": protocol["rationale"],
            "objectives": {
                "primary": protocol["objectives"]["primary"],
                "secondary": protocol["objectives"]["secondary"]
            },
            "study_design": protocol["study_design"],
            "population": {
                "inclusion_criteria": protocol["population"]["inclusion_criteria"],
                "exclusion_criteria": protocol["population"]["exclusion_criteria"]
            },
            "intervention": protocol["intervention"],
            "endpoints": {
                "primary": protocol["endpoints"]["primary"],
                "secondary": protocol["endpoints"]["secondary"]
            },
            "safety_monitoring": protocol["safety_monitoring"],
            "statistical_analysis": protocol["statistical_analysis"],
            "ethical_considerations": protocol["ethical_considerations"]
        }
    }

    path = f"/tmp/{filename}"
    with open(path, "w", encoding="utf-8") as f:
        json.dump(json_obj, f, indent=2, ensure_ascii=False)

    return path
